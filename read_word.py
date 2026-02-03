from docx import Document

doc = Document('docs/Bonjuk SSS - FAQ - Q&A v30122025 (1).docx')

print("=== WORD DOSYASI TAM İÇERİK ===\n")
print(f"Toplam {len(doc.paragraphs)} paragraf\n")

for i, para in enumerate(doc.paragraphs, 1):
    if para.text.strip():
        print(f"{i}. {para.text}")
        print()

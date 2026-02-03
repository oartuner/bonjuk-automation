import openpyxl
from docx import Document

# Excel dosyasını oku
wb = openpyxl.load_workbook('docs/2026 Bonjuk.xlsx')
print("=== EXCEL SHEETS ===")
for sheet_name in wb.sheetnames:
    print(f"\nSheet: {sheet_name}")
    sheet = wb[sheet_name]
    print(f"Rows: {sheet.max_row}, Columns: {sheet.max_column}")
    
    # İlk 10 satırı göster
    for i, row in enumerate(sheet.iter_rows(values_only=True), 1):
        if i > 10:
            break
        print(f"Row {i}: {row}")

print("\n\n=== WORD DOCUMENT ===")
# Word dosyasını oku
doc = Document('docs/Bonjuk SSS - FAQ - Q&A v30122025 (1).docx')
print(f"Total paragraphs: {len(doc.paragraphs)}")
print("\nFirst 20 paragraphs:")
for i, para in enumerate(doc.paragraphs[:20], 1):
    if para.text.strip():
        print(f"{i}. {para.text[:100]}...")

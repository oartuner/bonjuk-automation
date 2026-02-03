import openpyxl
from docx import Document
import json

# Excel dosyasını oku - Etkinlik Takvimi
wb = openpyxl.load_workbook('docs/2026 Bonjuk.xlsx')
sheet = wb['2026 Program']

print("=== 2026 ETKİNLİK TAKVİMİ ===\n")
events = []
for row in sheet.iter_rows(min_row=2, values_only=True):
    if row[0]:  # Tarih varsa
        date_range = str(row[0]).strip()
        event_name = str(row[1]).strip() if row[1] else "TBA"
        print(f"{date_range} -> {event_name}")
        events.append({"date_range": date_range, "event_name": event_name})

print(f"\nToplam {len(events)} etkinlik bulundu.")

# Word dosyasını oku - Email Şablonları
print("\n\n=== EMAIL ŞABLONLARI (WORD) ===\n")
doc = Document('docs/Bonjuk SSS - FAQ - Q&A v30122025 (1).docx')

# Tüm metni birleştir
full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# Şablon başlıklarını ara
templates_found = []
keywords = ["Sevgili", "Dear", "Merhaba", "Hello", "Rezervasyon", "Confirmation"]

current_template = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        # Yeni şablon başlangıcı olabilir
        if any(kw in text for kw in keywords) and len(text) < 100:
            if current_template:
                templates_found.append("\n".join(current_template))
                current_template = []
        current_template.append(text)

if current_template:
    templates_found.append("\n".join(current_template))

print(f"Toplam {len(templates_found)} potansiyel şablon bulundu.\n")
for i, template in enumerate(templates_found[:5], 1):  # İlk 5'ini göster
    print(f"\n--- Şablon {i} ---")
    print(template[:300] + "..." if len(template) > 300 else template)
    print()
